import fitz
from docx import Document
import re
import os

# PDF file name
pdf_file = "I Decided to Live as Me °.pdf"

# Open PDF
pdf = fitz.open(pdf_file)

# Get PDF name without .pdf
pdf_name = os.path.splitext(os.path.basename(pdf_file))[0]

# Output file name
output_file = f"{pdf_name}_Highlights.docx"

# Create Word document
doc = Document()

# Add title
doc.add_heading(f"{pdf_name} - Extracted Highlights", level=1)

# Avoid duplicates
seen_text = set()

# Function to clean formatting
def clean_text(text):

    # Replace line breaks with spaces
    text = text.replace("\n", " ")

    # Remove multiple spaces
    text = re.sub(r'\s+', ' ', text)

    return text.strip()

# Loop through pages
for page in pdf:

    annotations = page.annots()

    if annotations:

        for annot in annotations:

            # Only highlights
            if annot.type[1] == "Highlight":

                # Extract highlighted text
                highlighted_text = page.get_textbox(annot.rect)

                # Clean text
                highlighted_text = clean_text(highlighted_text)

                # Skip duplicates/empty
                if not highlighted_text or highlighted_text in seen_text:
                    continue

                seen_text.add(highlighted_text)

                # Heading detection
                if highlighted_text.isupper():

                    doc.add_heading(highlighted_text, level=2)

                else:

                    paragraph = doc.add_paragraph()
                    paragraph.add_run(highlighted_text)

# Save document
doc.save(output_file)

print(f"Document created: {output_file}")